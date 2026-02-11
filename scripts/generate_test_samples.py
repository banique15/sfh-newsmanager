
import os
import random
from reportlab.pdfgen import canvas
from docx import Document
from PIL import Image, ImageDraw, ImageFont

# Ensure directory exists
OUTPUT_DIR = "tests/samples"
os.makedirs(OUTPUT_DIR, exist_ok=True)

def generate_pdf(filename, pages=1):
    """Generate a PDF file with specified number of pages."""
    path = os.path.join(OUTPUT_DIR, filename)
    c = canvas.Canvas(path)
    
    for i in range(pages):
        c.drawString(100, 750, f"Page {i+1} of {pages}")
        c.drawString(100, 730, "This is a sample PDF generated for testing the SFH News Manager bot.")
        c.drawString(100, 710, "Use this file to verify PDF text extraction.")
        
        # Add some Lorem Ipsum content
        y = 690
        for _ in range(20): # Lines of text
            text = "Lorem ipsum dolor sit amet, consectetur adipiscing elit. " * 3
            c.drawString(100, y, text[:80])
            y -= 20
            
        c.showPage()
        
    c.save()
    print(f"Generated PDF: {path} ({pages} pages)")

def generate_docx(filename, paragraphs=5):
    """Generate a Word document."""
    path = os.path.join(OUTPUT_DIR, filename)
    doc = Document()
    doc.add_heading('Sample Word Document', 0)
    
    doc.add_paragraph("This is a sample document for testing the SFH News Manager bot.")
    
    for i in range(paragraphs):
        doc.add_paragraph(
            f"Paragraph {i+1}: Lorem ipsum dolor sit amet, consectetur adipiscing elit. "
            "Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua."
        )
        
    doc.save(path)
    print(f"Generated DOCX: {path}")

def generate_image(filename, size=(800, 600), color="blue"):
    """Generate a simple image."""
    path = os.path.join(OUTPUT_DIR, filename)
    img = Image.new('RGB', size, color=color)
    d = ImageDraw.Draw(img)
    
    # Add text to image
    try:
        # Try to load a font, fallback to default
        font = ImageFont.truetype("arial.ttf", 40)
    except IOError:
        font = ImageFont.load_default()
        
    d.text((50, 50), "Sample Image for Testing", fill=(255, 255, 255), font=font)
    d.text((50, 150), f"Size: {size[0]}x{size[1]}", fill=(255, 255, 255), font=font)
    
    # Draw some shapes
    d.rectangle([200, 200, 400, 400], outline="white", width=5)
    d.ellipse([450, 200, 650, 400], fill="red")
    
    img.save(path)
    print(f"Generated Image: {path}")

if __name__ == "__main__":
    print(f"Generating samples in {OUTPUT_DIR}...")
    
    # 1. Standard Files
    generate_pdf("sample_article.pdf", pages=2)
    generate_docx("sample_notes.docx", paragraphs=3)
    generate_image("sample_image.jpg", color="darkblue")
    
    # 2. Limit Testing Files
    # Large PDF (50 pages) - Simulates a report
    generate_pdf("large_report.pdf", pages=50)
    
    # Large DOCX (100 paragraphs)
    generate_docx("large_document.docx", paragraphs=100)
    
    # Large Image (4K resolution)
    generate_image("large_high_res.png", size=(3840, 2160), color="purple")
    
    print("Done!")
